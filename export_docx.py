"""
急诊抢救记录 DOCX 导出工具
----------------------------
用法：
  1. 从HTML工具中导出JSON数据文件
  2. 运行本脚本: python export_docx.py <json文件路径>
  3. 生成正式的DOCX格式抢救记录文档

依赖: pip install python-docx
"""

import json
import sys
import os
from datetime import datetime

def create_docx(data, output_path):
    """根据JSON数据生成DOCX文档"""
    from docx import Document
    from docx.shared import Pt, Cm, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn

    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # ===== 标题 =====
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('急 诊 抢 救 记 录')
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    # 分隔线
    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sep.add_run('━' * 30)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(128, 128, 128)

    # ===== 基本信息 =====
    patient = data.get('patient', {})
    doc.add_heading('基本信息', level=2)

    info_table = doc.add_table(rows=4, cols=4, style='Table Grid')
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    info_data = [
        ['姓名', patient.get('name', '________'), '性别', patient.get('gender', '')],
        ['年龄', patient.get('age', ''), '科室', patient.get('dept', '急诊科')],
        ['就诊号', patient.get('id', ''), '抢救日期', patient.get('date', '')],
        ['主诉', patient.get('chief', ''), '初步诊断', patient.get('diagnosis', '')],
    ]

    for i, row_data in enumerate(info_data):
        row = info_table.rows[i]
        for j, text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(text) if text else '________')
            run.font.size = Pt(10)
            if j % 2 == 0:  # Label columns
                run.font.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            # Handle merged cells for main complaint and diagnosis
            if i == 3 and j == 0:
                # Merge last row cells 1-3
                row.cells[1].merge(row.cells[3])

    # 合并主诉和诊断行
    info_table.cell(3, 0).merge(info_table.cell(3, 0))  # Already in correct cell

    # 重新处理最后一行
    row4 = info_table.rows[3]
    # Clear and reset
    for k in range(4):
        row4.cells[k].text = ''

    row4.cells[0].text = ''
    p0 = row4.cells[0].paragraphs[0]
    r0 = p0.add_run('主诉')
    r0.font.size = Pt(10)
    r0.font.bold = True
    p0.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Merge cells 1-3 for 主诉
    row4.cells[1].merge(row4.cells[3])
    row4.cells[1].text = ''
    p1 = row4.cells[1].paragraphs[0]
    r1 = p1.add_run(str(patient.get('chief', '________')))
    r1.font.size = Pt(10)

    # Add diagnosis row
    diag_row = info_table.add_row()
    diag_row.cells[0].text = ''
    p = diag_row.cells[0].paragraphs[0]
    r = p.add_run('初步诊断')
    r.font.size = Pt(10)
    r.font.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    diag_row.cells[1].merge(diag_row.cells[3])
    diag_row.cells[1].text = ''
    p = diag_row.cells[1].paragraphs[0]
    r = p.add_run(str(patient.get('diagnosis', '________')))
    r.font.size = Pt(10)

    doc.add_paragraph()  # spacing

    # ===== 抢救时间 =====
    rescue_start = data.get('rescueStartTime', '')
    rescue_end = data.get('rescueEndTime', '')

    if rescue_start or rescue_end:
        doc.add_heading('抢救时间', level=2)
        time_p = doc.add_paragraph()
        if rescue_start:
            time_p.add_run(f'抢救开始：{rescue_start[:5]}').font.size = Pt(10.5)
        if rescue_end:
            time_p.add_run(f'    抢救结束：{rescue_end[:5]}').font.size = Pt(10.5)
        if rescue_start and rescue_end:
            sh, sm = int(rescue_start[:2]), int(rescue_start[3:5])
            eh, em = int(rescue_end[:2]), int(rescue_end[3:5])
            diff = (eh*60+em) - (sh*60+sm) if (eh*60+em) >= (sh*60+sm) else (24*60 - sh*60 - sm + eh*60 + em)
            time_p.add_run(f'    抢救时长：约{diff}分钟').font.size = Pt(10.5)

    doc.add_paragraph()

    # ===== 抢救经过 =====
    items = data.get('items', [])
    if items:
        doc.add_heading('抢救经过', level=2)

        # Sort items by time
        sorted_items = sorted(items, key=lambda x: x.get('time', ''))

        type_labels = {'med': '💊 用药', 'measure': '🩺 措施', 'tube': '🩸 管道'}

        for item in sorted_items:
            p = doc.add_paragraph()
            t = item.get('time', '')[:5] if item.get('time') else '--:--'
            label = type_labels.get(item.get('type', ''), '')
            name = item.get('name', '')
            detail = item.get('detail', '')

            run_time = p.add_run(f'{t}  ')
            run_time.font.size = Pt(10.5)
            run_time.font.bold = True

            run_label = p.add_run(f'{label}  ')
            run_label.font.size = Pt(10.5)
            run_label.font.color.rgb = RGBColor(80, 80, 80)

            run_name = p.add_run(name)
            run_name.font.size = Pt(10.5)

            if detail and detail.strip():
                run_detail = p.add_run(f'（{detail}）')
                run_detail.font.size = Pt(9)
                run_detail.font.color.rgb = RGBColor(128, 128, 128)

        doc.add_paragraph()

    # ===== 分类汇总 =====
    meds = [i for i in items if i.get('type') == 'med']
    measures = [i for i in items if i.get('type') == 'measure']
    tubes = [i for i in items if i.get('type') == 'tube']

    if meds:
        doc.add_heading('用药情况汇总', level=2)
        for m in meds:
            p = doc.add_paragraph(style='List Bullet')
            t = m.get('time', '')[:5] if m.get('time') else '--:--'
            p.add_run(f'{t}  {m["name"]}').font.size = Pt(10.5)
            if m.get('detail'):
                p.add_run(f'  {m["detail"]}').font.size = Pt(9)

    if measures:
        doc.add_heading('抢救措施汇总', level=2)
        for m in measures:
            p = doc.add_paragraph(style='List Bullet')
            t = m.get('time', '')[:5] if m.get('time') else '--:--'
            p.add_run(f'{t}  {m["name"]}').font.size = Pt(10.5)
            if m.get('detail'):
                p.add_run(f'  {m["detail"]}').font.size = Pt(9)

    if tubes:
        doc.add_heading('管道通路汇总', level=2)
        for t in tubes:
            p = doc.add_paragraph(style='List Bullet')
            tt = t.get('time', '')[:5] if t.get('time') else '--:--'
            p.add_run(f'{tt}  {t["name"]}').font.size = Pt(10.5)
            if t.get('detail'):
                p.add_run(f'  {t["detail"]}').font.size = Pt(9)

    # ===== 抢救结果 =====
    doc.add_heading('抢救结果', level=2)
    outcomes = [
        '□ 病情稳定，转入病房/ICU',
        '□ 抢救成功，留观',
        '□ 自动出院',
        '□ 死亡',
        '□ 其他：________',
    ]
    for o in outcomes:
        p = doc.add_paragraph(o)
        p.paragraph_format.space_after = Pt(2)

    doc.add_paragraph()

    # ===== 签名栏 =====
    doc.add_paragraph()
    sign_p = doc.add_paragraph()
    sign_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sign_p.add_run('记录医师：________').font.size = Pt(10.5)

    time_p = doc.add_paragraph()
    time_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    time_p.add_run(f'记录时间：{data.get("generatedAt", datetime.now().strftime("%Y-%m-%d %H:%M:%S"))}').font.size = Pt(10.5)

    # 保存
    doc.save(output_path)
    print(f'✅ DOCX文档已生成: {output_path}')

def main():
    if len(sys.argv) < 2:
        print('用法: python export_docx.py <json文件路径>')
        print('示例: python export_docx.py 急诊抢救记录_张三_2026-06-06.json')
        sys.exit(1)

    json_path = sys.argv[1]
    if not os.path.exists(json_path):
        print(f'❌ 文件不存在: {json_path}')
        sys.exit(1)

    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    # 生成输出路径
    base = os.path.splitext(json_path)[0]
    output_path = f'{base}.docx'

    create_docx(data, output_path)


if __name__ == '__main__':
    main()
